"""Genera los dos documentos Word para el cliente, en la carpeta de Descargas.

Por qué un script y no dos `.md` más en el repositorio: ya hay dos documentos de requisitos
(`INSTALACION_CLIENTE.md`, ejecutivo, y `REQUISITOS_INSTALACION.md`, detallado) y añadir un
tercero es fabricar la siguiente contradicción. Aquí las cifras se escriben UNA vez y el
`.docx` se regenera; el documento que **prevalece** sigue siendo `REQUISITOS_INSTALACION.md`.

El primero está acotado a **una página** — comprobado con Word (`ComputeStatistics`) y mirando
el PDF renderizado, no estimado. Si se le añade contenido hay que volver a comprobarlo.

    python scripts/generar_docs_cliente.py
    python scripts/generar_docs_cliente.py --salida /otra/carpeta
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    raise SystemExit("Falta python-docx.  pip install python-docx")

REPO = Path(__file__).resolve().parent.parent
AZUL = RGBColor(0x1F, 0x3B, 0x57)
GRIS = RGBColor(0x55, 0x5F, 0x6B)
ROJO = RGBColor(0xA3, 0x2A, 0x2A)
VERDE = RGBColor(0x1E, 0x6B, 0x3A)


# --------------------------------------------------------------------------- utilidades
def _base(doc: Document, pt: float, margen: float) -> None:
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(pt)
    st.paragraph_format.space_after = Pt(3)
    st.paragraph_format.space_before = Pt(0)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(margen)
        s.left_margin = s.right_margin = Cm(margen)


def _p(doc, texto="", *, pt=None, negrita=False, color=None, antes=0, despues=3,
       cursiva=False, izq=0):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(antes)
    par.paragraph_format.space_after = Pt(despues)
    if izq:
        par.paragraph_format.left_indent = Cm(izq)
    if texto:
        _run(par, texto, pt=pt, negrita=negrita, color=color, cursiva=cursiva)
    return par


def _run(par, texto, *, pt=None, negrita=False, color=None, cursiva=False):
    r = par.add_run(texto)
    r.bold = negrita
    r.italic = cursiva
    if pt:
        r.font.size = Pt(pt)
    if color is not None:
        r.font.color.rgb = color
    return r


def _rico(par, trozos, *, pt=None):
    """trozos: lista de (texto, negrita) o (texto, negrita, color)."""
    for t in trozos:
        _run(par, t[0], pt=pt, negrita=t[1], color=(t[2] if len(t) > 2 else None))


def _titulo(doc, texto, *, pt=15):
    par = _p(doc, despues=2)
    _run(par, texto, pt=pt, negrita=True, color=AZUL)
    return par


def _seccion(doc, texto, *, pt=10.5, antes=8):
    par = _p(doc, antes=antes, despues=2)
    _run(par, texto, pt=pt, negrita=True, color=AZUL)
    return par


def _sombra(celda, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    celda._tc.get_or_add_tcPr().append(el)


def _tabla(doc, cabecera, filas, anchos, *, pt=8.5, fondo="1F3B57"):
    t = doc.add_table(rows=1, cols=len(cabecera))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, txt in enumerate(cabecera):
        c = t.rows[0].cells[i]
        c.text = ""
        _sombra(c, fondo)
        par = c.paragraphs[0]
        par.paragraph_format.space_after = Pt(1)
        par.paragraph_format.space_before = Pt(1)
        _run(par, txt, pt=pt, negrita=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for fila in filas:
        celdas = t.add_row().cells
        for i in range(len(celdas)):
            c = celdas[i]
            c.text = ""
            par = c.paragraphs[0]
            par.paragraph_format.space_after = Pt(1)
            par.paragraph_format.space_before = Pt(1)
            dato = fila[i]
            if isinstance(dato, list):
                _rico(par, dato, pt=pt)
            else:
                _run(par, str(dato), pt=pt)
    for fila in t.rows:
        for i, c in enumerate(fila.cells):
            c.width = Cm(anchos[i])
    return t


def _vineta(doc, trozos, *, pt=9, izq=0.45):
    par = _p(doc, despues=2, izq=izq)
    _run(par, "•  ", pt=pt, negrita=True, color=AZUL)
    _rico(par, trozos, pt=pt)
    return par


def _pie(doc, texto, *, pt=7.5):
    par = _p(doc, antes=8, despues=0)
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(par, texto, pt=pt, color=GRIS, cursiva=True)
    return par


# --------------------------------------------------------------- 1) servidor, UNA página
def doc_servidor(destino: Path) -> Path:
    doc = Document()
    _base(doc, 9, 1.45)

    _titulo(doc, "Servidor y software para el lector de incapacidades")
    par = _p(doc, despues=5)
    _rico(par, [
        ("En una línea: ", True),
        ("un servidor Linux con ", False),
        ("Docker", True),
        (", sin salida a internet, ", False),
        ("8 núcleos, 16 GB de RAM y 500 GB de disco", True),
        (". Con 4 núcleos y 250 GB funciona, pero el disco solo cubre unos 3 años.", False),
    ], pt=9)

    _seccion(doc, "1. Software a instalar — solo hace falta Docker", antes=4)
    _tabla(doc,
           ["Componente", "¿Obligatorio?", "Espacio"],
           [
               [[("Docker Engine", True), (" ≥ 23.0 + plugin ", False), ("Compose v2", True)],
                [("SÍ", True, VERDE)], "~0,5 GB"],
               ["Imagen de la aplicación (se construye una vez)",
                [("SÍ", True, VERDE)], "~1,1 GB"],
               [[("MySQL 8", True), (" — solo para pruebas", False)],
                "NO en producción: allí se apunta a la base de datos del ERP", "~0,6 GB"],
               [[("Ollama", True), (" + modelos de IA local (permisos manuscritos)", False)],
                "NO — sin él, esos casos van a revisión manual", "~14,5 GB"],
           ],
           [7.4, 7.6, 2.6])
    par = _p(doc, antes=3, despues=4)
    _rico(par, [
        ("No hace falta instalar Python, ni Poppler, ni Tesseract, ni ningún servicio de OCR: ",
         True),
        ("todo viaja dentro de la imagen de Docker.", False),
    ], pt=8.5)

    _seccion(doc, "2. Hardware")
    _tabla(doc,
           ["", "Mínimo", "Recomendado", "Con IA local"],
           [
               ["Núcleos físicos (x86-64)", "4", [("8", True)], "16"],
               ["RAM", "8 GB (con el tope de píxeles puesto)", [("16 GB", True)], "32 GB"],
               ["Disco", "250 GB SSD → unos 3 años",
                [("500 GB SSD NVMe → 5 años", True)], "1 TB SSD NVMe"],
               ["Sistema operativo", "Linux x86-64 (Ubuntu Server LTS / RHEL)", "igual", "igual"],
               ["Tarjeta gráfica", [("no hace falta", True)], "no hace falta",
                "opcional (solo acelera la IA)"],
           ],
           [4.0, 4.6, 4.6, 4.4])
    par = _p(doc, antes=3, despues=4)
    _rico(par, [
        ("El volumen de Gruppo no necesita CPU: necesita RAM y disco. ", True),
        ("Medido sobre los 31 documentos reales: ", False),
        ("10 s de CPU por documento", True),
        (" (7.000 al mes = 19,4 horas de CPU mensuales) y ", False),
        ("379 KB por documento", True),
        (" (unos 65 GB al año).", False),
    ], pt=8.5)

    _seccion(doc, "3. Lo que el sistema NO necesita")
    _vineta(doc, [("Internet en funcionamiento.", True),
                  (" Ni los documentos ni los datos salen del servidor (datos de salud,"
                   " Ley 1581). Internet solo se usa UNA vez, para bajar las imágenes; si el"
                   " servidor está aislado, el traslado se hace con un archivo.", False)], pt=8.5)
    _vineta(doc, [("Tarjeta gráfica.", True), (" Todo corre en CPU.", False)], pt=8.5)
    _vineta(doc, [("Licencias.", True),
                  (" Todo el software es libre y no hay APIs de pago.", False)], pt=8.5)

    _seccion(doc, "4. Instalación y comprobación")
    for linea in ("docker compose up -d --build",
                  "curl http://localhost:8000/api/health     → debe responder  "
                  "{\"status\":\"ok\"}"):
        par = _p(doc, despues=1, izq=0.4)
        r = _run(par, linea, pt=8.5)
        r.font.name = "Consolas"
    par = _p(doc, antes=3, despues=4)
    _rico(par, [
        ("La aplicación queda en ", False), ("http://localhost:8000", True),
        (", publicada ", False), ("solo en el propio servidor", True),
        (" (los documentos tienen datos personales). Para abrirla a la red interna hay que"
         " añadir un proxy con cifrado y usuarios: hoy no los trae.", False),
    ], pt=8.5)

    _seccion(doc, "5. Tres avisos honestos antes de comprar")
    _vineta(doc, [("Las cifras se midieron en un portátil de 15 W", True),
                  (" y el mismo documento varió hasta ×2,86 entre dos pasadas: son un orden de"
                   " magnitud, no un compromiso. Se entrega el script para repetir la medición"
                   " en el servidor real antes de cerrar la compra.", False)], pt=8.5)
    _vineta(doc, [("La RAM es el dato que hay que mirar.", True),
                  (" Un PDF del corpus llegó a un pico de 7,6 GB con los valores por defecto."
                   " Hay que poner ", False), ("OCR_MAX_PIXELS=8000000", True),
                  (" para que 8 GB alcancen; está medido sobre 5 documentos y hay que"
                   " confirmarlo en el servidor.", False)], pt=8.5)
    _vineta(doc, [("Cada corrida procesa como máximo 500 casos", True),
                  (", así que un día pico (unos 875 documentos) necesita dos corridas hasta que"
                   " subamos ese tope. Es un arreglo de software, no de hardware.", False)],
            pt=8.5)

    _seccion(doc, "6. Lo que necesitamos del cliente para instalar")
    par = _p(doc, despues=2, izq=0.4)
    _rico(par, [
        ("Servidor con lo de arriba y Docker como servicio", True),
        (" (que arranque solo tras un reinicio) · ", False),
        ("acceso a la base de datos del ERP", True),
        (" (servidor, puerto, usuario, clave) · ", False),
        ("carpeta compartida", True),
        (" donde recepción deja los documentos · ", False),
        ("decisión Linux o Windows", True),
        (" (recomendamos Linux) · ", False),
        ("plazo legal de conservación", True),
        (" de los soportes (decide entre 250 GB y 1 TB) · ", False),
        ("exclusión del antivirus", True),
        (" sobre la carpeta de entrada, con escaneo diario en modo alertar · ", False),
        ("cifrado del disco", True),
        (" antes de escribir el primer documento · ", False),
        ("¿el servidor tendrá TPM?", True),
        (" Sin TPM no se puede tener cifrado del disco y reinicio desatendido a la vez, y eso"
         " es un requisito de compra, no de configuración.", False),
    ], pt=8.5)

    _pie(doc, "Documento resumen. El detalle completo (comandos del traslado a un equipo "
              "aislado, matriz Windows vs. Linux y la aritmética de CPU, RAM y disco) está en "
              "REQUISITOS_INSTALACION.md, que es el documento que PREVALECE si alguna cifra "
              "no coincide.")

    ruta = destino / "incapacidad-ocr - Servidor y software necesario.docx"
    doc.save(ruta)
    return ruta


# ------------------------------------------------- 2) cambios, pruebas y pendientes
def doc_estado(destino: Path) -> Path:
    doc = Document()
    _base(doc, 9.5, 1.6)

    _titulo(doc, "Lector de incapacidades — qué cambió, qué se probó y qué falta")
    par = _p(doc, despues=6)
    _rico(par, [
        ("Estado: ", True),
        ("listo para que Gruppo lo pruebe con sus propios documentos. ", False),
        ("Producción todavía no", True),
        (" — falta el servidor y los datos que solo el cliente tiene.", False),
    ], pt=9.5)

    # ---------------------------------------------------------------- 1. qué cambió
    _seccion(doc, "1. Qué se hizo", pt=11, antes=4)
    _tabla(doc,
           ["Cambio", "Qué resuelve", "Efecto medido"],
           [
               [[("Detección de documentos adulterados", True)],
                "Se cargó el catálogo CIE-10 completo (14.484 códigos, de fuente pública, "
                "incluido en el sistema y sin depender de internet). Antes, con un catálogo "
                "hecho a mano, no se podía afirmar que un código no existiera",
                [("de 2 a 4", True), (" de 9 trámites adulterados, con ", False),
                 ("0 falsas alarmas", True)]],
               [[("Validación de fechas y días", True)],
                "Motor de 17 reglas que compara lo IMPRESO en el papel —nunca lo que el sistema "
                "calcula— y explica por qué. La severidad y los umbrales se cambian desde la "
                "base de datos, sin volver a instalar nada",
                "Encontró un desfase de 30 días que Gruppo había marcado por otro motivo"],
               [[("Días en letras y en números", True)],
                "Lee la duración escrita como 2, como DOS y como las dos a la vez "
                "(DOS (2) DÍAS, 30 (TREINTA), 14 - CATORCE). Cuando hay palabra y dígito manda "
                "el dígito, y el desacuerdo se registra",
                "13 formas reales del corpus, y los falsos positivos que se rechazan"],
               [[("Organización de la carpeta de ingesta", True)],
                "Tres zonas que se leen en orden de flujo (1_entrada, 2_revisar, 3_archivo), con "
                "una subcarpeta por MOTIVO: la carpeta dice qué hay que hacer",
                "Cada archivo termina en UNA sola zona: «¿dónde quedó?» tiene respuesta única"],
               [[("Prueba repetible", True)],
                "Botón «Reiniciar prueba»: devuelve los 31 documentos a la entrada y limpia la "
                "bandeja. Base de datos de prueba en Docker, con catálogos calcados de los "
                "documentos reales",
                "31 documentos restaurados y bandeja a 0, tantas veces como se quiera"],
               [[("La aclaración del rojo", True), (" de la tabla de motivos", False)],
                "Se aplicó la regla que confirmó Diana: el rojo marca que el documento está mal "
                "y que la razón es la de la fila inmediatamente anterior",
                [("3 documentos", True),
                 (" que estaban «sin motivo» ya se pueden evaluar", False)]],
               [[("El proyecto se instala solo", True)],
                "Se encontró que la base de datos NO podía inicializarse en una máquina nueva: "
                "una instrucción válida en otro motor abortaba el arranque y el contenedor "
                "moría. Corregido y comprobado con una base vacía",
                "«git clone» más un comando y queda operativo, con el catálogo ya cargado"],
           ],
           [4.2, 8.0, 5.2])

    # ---------------------------------------------------------------- 2. qué se probó
    _seccion(doc, "2. Qué se probó, y con qué resultado", pt=11)
    _tabla(doc,
           ["Prueba", "Resultado"],
           [
               [[("Los 31 documentos reales de Gruppo", True),
                 (" (15 marcados como adulterados, 16 legítimos) procesados en lote", False)],
                [("27 trámites", True),
                 (" — la llave es la cédula, así que los documentos de una misma persona son "
                  "un solo trámite. Ninguno con error, ninguno mal nombrado", False)]],
               ["Detección de adulteración",
                [("4 de 9", True), (" trámites adulterados detectados y ", False),
                 ("0 falsos positivos", True),
                 (" sobre los 13 trámites legítimos. Los 5 casos en cuarentena —los de "
                  "etiqueta en disputa— no cuentan ni a favor ni en contra", False)]],
               ["Exactitud de lectura sobre los 8 documentos de la primera muestra",
                [("82 %", True), (" de los campos núcleo en el entorno de producción "
                 "(Python 3.12); 76 % en el de desarrollo, que arrastra una versión antigua del "
                 "motor de OCR. Pendiente de volver a verificarlo en el servidor definitivo",
                 False)]],
               ["Baterías de pruebas automáticas",
                [("8 de 8 en verde", True),
                 (", sin necesitar los documentos ni la base de datos", False)]],
               ["Instalación desde cero",
                [("Verificada dos veces", True),
                 (": una copia limpia del repositorio (232 archivos) pasa las 8 baterías, y una "
                  "base de datos vacía queda con las 14 tablas y los 14.484 diagnósticos sin "
                  "ejecutar nada más", False)]],
               ["Que no se marque un documento bueno",
                [("Es el criterio que manda", True),
                 (": un falso positivo se paga bajando la severidad o exigiendo más evidencia, ",
                  False),
                 ("nunca moviendo el umbral hasta acertar", True),
                 (" en 31 documentos. Con 7.000 casos al mes, un sistema que desconfía de los "
                  "documentos buenos hace que el auxiliar deje de mirar las alertas", False)]],
           ],
           [5.4, 12.0])
    par = _p(doc, antes=3, despues=4)
    _rico(par, [
        ("Aviso sobre estas cifras: ", True),
        ("31 documentos ", False),
        ("no sirven para prometer una precisión", True),
        (". Sirven para descartar controles malos —ya descartaron varios— y como prueba de que "
         "nada se rompe. Para dar un número defendible hace falta lo de la sección 3.", False),
    ], pt=8.5)

    doc.add_page_break()

    # ------------------------------------------------------- 3. pendiente de Diana
    _seccion(doc, "3. Lo que queda pendiente del lado de Diana", pt=11, antes=0)
    par = _p(doc, despues=5)
    _rico(par, [("Primero lo urgente. Estas seis mueven más que todo el resto junto.", True)],
          pt=9.5)
    _tabla(doc,
           ["#", "Qué se necesita", "Por qué urge"],
           [
               [[("D1", True, ROJO)],
                [("Cuando el papel imprime los días ", False), ("y", True),
                 (" la fecha de fin y no cuadran, ¿qué manda?", False)],
                [("Es la más urgente. ", True, ROJO),
                 ("Hoy el sistema decide en silencio que manda «inicio + días» y ", False),
                 ("reescribe", True),
                 (" la fecha de fin. Si la respuesta es la contraria, hay un dato equivocado "
                  "entrando a nómina hoy mismo", False)]],
               [[("A1", True, ROJO)],
                "Datos de conexión a la base de datos real del ERP, con lectura sobre los "
                "listados y escritura sobre las dos tablas del middleware",
                [("Es el interruptor que convierte la demostración en piloto.", True),
                 (" Hoy todo escribe en una base de juguete de unas 30 filas", False)]],
               [[("A2", True, ROJO)],
                "Listado de personal vigente: código interno, cédula, nombre completo y EPS. "
                "Puede venir tal cual sale del ERP",
                [("Sin esto ", False), ("ninguna", True),
                 (" cédula resuelve y los 7.000 casos al mes caen a «datos por revisar», que es "
                  "exactamente el trabajo manual que el proyecto elimina", False)]],
               [[("A3", True)],
                "Listado de diagnósticos que usa el ERP (código y descripción), en Excel o CSV",
                [("Permite responder «¿está en el catálogo de Gruppo?» en vez de «¿existe en la "
                  "CIE-10 pública?». Es el motivo ", False),
                 ("más repetido", True),
                 (" de la tabla de motivos, y hoy no se puede comprobar", False)]],
               [[("C4", True)],
                "Cuál es la etiqueta correcta de los 5 documentos contradictorios: dos parejas "
                "de archivos idénticos entregados como falso Y como legítimo, y uno que "
                "comparte cédula con un legítimo",
                [("Mientras no se resuelva, ", False), ("ningún porcentaje es defendible", True),
                 (": en una pareja contradictoria el motor falla 1 de 2 responda lo que "
                  "responda. Resolverlo devuelve 5 documentos —el 16 % del corpus— a toda "
                  "medición", False)]],
               [[("D9", True)],
                "Aprobar las 10 filas de requisitos documentales internos por tipo de "
                "ausentismo (por ejemplo, maternidad = incapacidad + historia clínica + "
                "nacido vivo o registro civil)",
                [("Decide si un caso sale «incompleto». Con 7.000 trámites al mes, ", False),
                 ("un requisito de más produce miles de alertas falsas", True),
                 (" y el equipo deja de mirarlas", False)]],
           ],
           [1.1, 7.1, 9.2])

    _seccion(doc, "El resto, agrupado (todo el detalle está en PENDIENTES.md)")
    _vineta(doc, [("Volcado del ERP, 6 listados (A1 a A6). ", True),
                  ("Personal, diagnósticos, EPS y ARL, estados de recepción, niveles de "
                   "incapacidad y el histórico de ausentismos ya radicados — este último "
                   "sin nombres ni cédulas: basta un identificador anónimo. El histórico "
                   "habilita 3 de las 17 reglas de fechas, hoy escritas, probadas y apagadas "
                   "por falta de acceso.", False)])
    _vineta(doc, [("Los números que deciden qué servidor comprar, 12 preguntas (B1 a B12). ",
                   True),
                  ("Van en un solo correo. Las tres que más mueven: ¿los 7.000 al mes son "
                   "trámites o archivos?, ¿cuántos años hay que conservar el original del "
                   "soporte? —decide entre un disco de 250 GB y uno de 1 TB— y ¿de cada 100 "
                   "radicadas, cuántas resultan adulteradas? Sin esta última, el umbral de "
                   "sospecha es una preferencia nuestra y no una decisión de negocio.", False)])
    _vineta(doc, [("Más documentos para poder dar cifras, 7 peticiones (C1 a C7). ", True),
                  ("100 o más legítimos exportados desde Word —el flujo que genera falsas "
                   "alarmas, y del que el corpus no tiene ni uno—, 20 o más adulterados por "
                   "señal, y 5 a 10 paquetes escaneados de 10 a 30 páginas. Y algo importante: "
                   "que los próximos lotes lleguen con el mismo patrón de nombre en las dos "
                   "clases, porque hoy un clasificador que solo lea el nombre del archivo "
                   "acierta el 100 % sin abrir el documento.", False)])
    _vineta(doc, [("Preguntas de operación y de norma, 13 respuestas cortas (D1 a D13). ", True),
                  ("Cada una se aplica cambiando un parámetro en la base de datos, sin volver a "
                   "instalar nada. Además de D1 y D9: ¿hay alguna EPS que en «Fecha Fin» imprima "
                   "el día de reintegro en vez del último día incapacitado?, ¿en cuántos días "
                   "hay que radicar para que se pague? y ¿existe ya el proceso del ERP que toma "
                   "las filas aprobadas y crea el ausentismo real? Esa última es el último "
                   "eslabón: si no existe, el auxiliar aprueba y no pasa nada en el ERP.",
                   False)])
    _vineta(doc, [("Para TI, no para Diana, 7 puntos (E1 a E7). ", True),
                  ("TPM, política de virtualización, respaldo local, quién administra los "
                   "contenedores, antivirus, cifrado del disco y dónde vive la base de datos "
                   "del ERP.", False)])

    # ------------------------------------------------------- 4. pendiente nuestro
    _seccion(doc, "4. Lo que falta de nuestro lado (no depende del cliente)")
    _vineta(doc, [("El lector toma el rótulo de la columna como si fuera el nombre del "
                   "diagnóstico", True),
                  (" —lee literalmente «CIE10»—, así que no hay texto que comparar. Es el "
                   "arreglo con más retorno: recupera 2 de los 5 casos que hoy no se detectan, "
                   "y es esfuerzo bajo.", False)])
    _vineta(doc, [("Registrar de cada dato si fue LEÍDO o CALCULADO", True),
                  (". Habilita 1 caso más y varias reglas de fechas que hoy están apagadas "
                   "porque no pueden distinguir el dato impreso del que el sistema deduce.",
                   False)])
    _vineta(doc, [("Subir el tope de 500 casos por corrida", True),
                  (". Un día pico son unos 875 documentos, así que hoy necesitaría dos "
                   "corridas.", False)])
    _vineta(doc, [("Los permisos manuscritos", True),
                  (" los lee mal el motor rápido. El de IA local los lee mucho mejor, pero es "
                   "lento en CPU, y el recuadro remunerado / no remunerado ", False),
                  ("no se detecta de forma confiable con ninguno de los dos", True),
                  (": ese campo lo elige el auxiliar a mano. Es el comportamiento esperado, no "
                   "un defecto pendiente.", False)])

    _pie(doc, "Documento resumen. El inventario completo y verificado —45 peticiones al "
              "cliente, agrupadas en 5 bloques, más el trabajo pendiente propio con su "
              "impacto— está en PENDIENTES.md. Para probarlo: MANUAL_PRUEBA.md, 15 minutos, "
              "solo necesita Docker.")

    ruta = destino / "incapacidad-ocr - Cambios, pruebas y pendientes.docx"
    doc.save(ruta)
    return ruta


def main() -> int:
    ap = argparse.ArgumentParser(description="Genera los documentos Word del cliente.")
    ap.add_argument("--salida", default=str(Path.home() / "Downloads"),
                    help="Carpeta de destino (por defecto, Descargas).")
    args = ap.parse_args()
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:  # noqa: BLE001
        pass

    destino = Path(args.salida)
    if not destino.is_dir():
        raise SystemExit(f"No existe la carpeta de salida: {destino}")
    for fn in (doc_servidor, doc_estado):
        ruta = fn(destino)
        print(f"generado: {ruta.name}  ({ruta.stat().st_size / 1024:.0f} KB)")
    print(f"\nen: {destino}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
